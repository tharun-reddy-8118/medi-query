from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Initialize Document
doc = Document()

# Define Styles
title_style = doc.styles['Title']
title_font = title_style.font
title_font.name = 'Calibri'
title_font.size = Pt(28)
title_font.color.rgb = RGBColor(34, 43, 69)

h1_style = doc.styles['Heading 1']
h1_font = h1_style.font
h1_font.name = 'Calibri'
h1_font.size = Pt(16)
h1_font.bold = True
h1_font.color.rgb = RGBColor(63, 81, 181)

h2_style = doc.styles['Heading 2']
h2_font = h2_style.font
h2_font.name = 'Calibri'
h2_font.size = Pt(14)
h2_font.bold = True
h2_font.color.rgb = RGBColor(34, 43, 69)

# Title
title = doc.add_paragraph('MediQuery: AI-Powered Data Analytics Platform', style='Title')
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('Platform Overview, Architecture, and Technical Specifications\n').alignment = WD_ALIGN_PARAGRAPH.CENTER

# 1. Executive Summary
doc.add_heading('1. Executive Summary', level=1)
doc.add_paragraph(
    "MediQuery is an advanced, AI-assisted data analytics and visualization platform designed to accelerate "
    "dashboard creation and data discovery. Unlike traditional platforms where end-users prompt AI constantly, "
    "MediQuery uses AI strategically during the initial data onboarding to generate a baseline dashboard. "
    "From there, Administrators take full control—manually building custom charts, defining complex metrics, "
    "and designing highly polished dashboards for end-users to consume and interact with. This approach guarantees "
    "extreme scalability, zero ongoing AI API costs for dashboard viewing, and absolute data privacy."
)

# 2. Data Connectivity & Integration
doc.add_heading('2. Data Connectivity', level=1)
doc.add_paragraph("MediQuery currently supports robust data integration pathways to connect to enterprise data:")
data_ways = [
    ("Direct File Upload", "Administrators can directly upload structured CSV, TSV, or Excel files. The backend seamlessly parses the file into the high-performance DuckDB analytical engine."),
    ("JSON Upload", "Structured JSON datasets can be flattened and integrated directly into the dashboard ecosystem."),
    ("Future/Extendable Connections", "Because the core engine is DuckDB, the backend can easily be extended to directly query external databases (e.g., PostgreSQL, MySQL, Snowflake) or read parquet files directly from S3 buckets using DuckDB's robust extension ecosystem.")
]
for t, d in data_ways:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(f"{t}: ").bold = True
    p.add_run(d)

# 3. Architecture & Workflow
doc.add_heading('3. System Architecture & Workflows', level=1)

doc.add_heading('System Diagram', level=2)
doc.add_paragraph("The following diagram outlines the overall architecture, distinguishing between the core execution engine and the LLM API layer.")
try:
    doc.add_picture('arch.png', width=Inches(6.0))
except Exception as e:
    doc.add_paragraph(f"[Image placeholder: arch.png not found - {str(e)}]")

doc.add_heading('How Charts are Generated', level=2)
chart_workflow = [
    ("Initial Discovery", "When data is uploaded, the backend generates a 'Data Dictionary' containing column names, types, and sample data."),
    ("Baseline Generation via AI", "This schema is securely sent to the Groq API (Llama 3.1). The AI responds with a JSON array outlining 4-6 foundational charts, mapping columns to X/Y axes and applying basic aggregations."),
    ("Admin Orchestration", "The administrator reviews the baseline, modifies colors, adds Calculated Fields, and manually creates any additional charts using the Drag-and-Drop UI. No AI is used for manual chart creation."),
    ("Execution", "The React frontend sends chart configurations to the FastAPI backend, which translates them into blazing-fast DuckDB SQL queries. DuckDB aggregates the data locally and returns only the summary values to the frontend to render the SVGs.")
]
for i, (t, d) in enumerate(chart_workflow, 1):
    p = doc.add_paragraph(style='List Number')
    p.add_run(f"{t}: ").bold = True
    p.add_run(d)

doc.add_heading('How "Ask AI" Works', level=2)
ask_ai = [
    ("Context Gathering", "When a user asks a question via the Chat interface, the frontend packages the query along with the schema and the metadata of the currently visible dashboard charts."),
    ("Querying the LLM", "The prompt is sent to the Groq Llama-3.1 model. The AI is instructed to act as a data analyst. It decides whether to generate a SQL query to answer a specific numerical question or simply provide a natural language explanation of the data."),
    ("Execution & Response", "If SQL is generated, the backend executes it against DuckDB and feeds the result back to the LLM to summarize, or the frontend directly displays the answer/chart to the user.")
]
for i, (t, d) in enumerate(ask_ai, 1):
    p = doc.add_paragraph(style='List Number')
    p.add_run(f"{t}: ").bold = True
    p.add_run(d)

# 4. Technologies Used
doc.add_heading('4. Technologies Used', level=1)
tech_stack = [
    ("Frontend", "React, Vite, Material-UI (MUI), Recharts (for SVGs), and React-Grid-Layout."),
    ("Backend", "FastAPI (Python) and Uvicorn for asynchronous, high-throughput API endpoints."),
    ("Database Engine", "DuckDB — An embedded analytical SQL database optimized for blazing-fast OLAP queries."),
    ("AI & NLP Engine", "Groq API utilizing the 'llama-3.1-8b-instant' model for initial schema parsing and chat functionality.")
]
for t, d in tech_stack:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(f"{t}: ").bold = True
    p.add_run(d)

# 5. Deployment & Scalability
doc.add_heading('5. Deployment Strategy', level=1)
doc.add_paragraph(
    "MediQuery is built with a modern, container-friendly architecture that can be deployed across various cloud providers (AWS, Azure, GCP) or on-premise infrastructure."
)
deployment = [
    ("Dockerization", "Both the frontend and backend can be containerized using Docker. The backend runs as a stateless container, while the frontend is served as static HTML/JS assets via Nginx."),
    ("Backend Scaling", "The FastAPI backend can be horizontally scaled across a Kubernetes cluster. Because DuckDB can read/write to shared persistent volumes or S3, multiple backend replicas can serve thousands of concurrent requests."),
    ("Zero Ongoing API Costs", "Because the AI is only used during initial data upload or explicit 'Ask AI' queries, daily end-user dashboard traffic does not trigger LLM API calls. This makes the platform incredibly cost-effective at scale.")
]
for t, d in deployment:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(f"{t}: ").bold = True
    p.add_run(d)

# 6. Groq API Limits & System Capabilities
doc.add_heading('6. Limits & System Capabilities', level=1)
doc.add_paragraph(
    "By restricting AI usage to the initial data onboarding phase and optional chat interactions, the platform gracefully bypasses typical AI limitations:"
)

groq_limits = [
    ("Llama 3.1 8B Output Limits (8,192 Tokens)", "The Llama-3.1-8b-instant model has a maximum output limit of 8,192 tokens. Because MediQuery only asks the AI to generate a *few* foundational charts initially, it comfortably stays well within this token limit."),
    ("Unlimited Administrator Dashboarding", "Because manual chart creation relies entirely on the local backend DuckDB instance and uses 0 LLM tokens, Administrators can add an unlimited number of widgets without hitting any Groq token caps or rate limits."),
    ("Groq Rate Limits Avoided", "Groq imposes Limits on Requests Per Minute (RPM) and Tokens Per Minute (TPM) based on pricing tiers. However, because end-users consuming the dashboard do not generate new charts via AI, millions of users can interact with the dashboard simultaneously without sending a single request to Groq.")
]
for t, d in groq_limits:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(f"{t}: ").bold = True
    p.add_run(d)

doc.save('MediQuery_Platform_Overview_v4.docx')
